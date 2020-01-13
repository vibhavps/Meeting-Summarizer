# -*- coding: utf-8 -*-
"""
Created on Wed Aug  7 19:25:35 2019

@author: nitin
"""

import nltk
from nltk.corpus import stopwords
import pandas as pd
from wordcloud import WordCloud
import matplotlib.pyplot as plt
import os
import re
 
from nltk.collocations import BigramAssocMeasures, BigramCollocationFinder
 
from operator import itemgetter
nltk.download('punkt') 
import docx 
from docx.shared import Inches
######################Summary################
doc = docx.Document() 
doc.add_heading('Meeting Notes - Global Executive Dashboard', 0) 
doc.add_heading('Summary', 1) 
df = pd.read_csv("transcript2.csv")
#df = open('transcript.txt').read().strip()
df.head()
df['article_text'][0]


from nltk.tokenize import sent_tokenize
sentences = []
for s in df['article_text']:
  sentences.append(sent_tokenize(s))

sentences = [y for x in sentences for y in x] # flatten list

sentences[:5]

#https://nlp.stanford.edu/projects/glove/

# Extract word vectors
os.getcwd()
word_embeddings = {}

f = open('glove.6B.100d.txt', encoding='utf-8')
for line in f:
    values = line.split()
    word = values[0]
    coefs = np.asarray(values[1:], dtype='float32')
    word_embeddings[word] = coefs
f.close()

len(word_embeddings)
#We now have word vectors for 400,000 different 
#terms stored in the dictionary – ‘word_embeddings’.



#Text Preprocessing

# remove punctuations, numbers and special characters
clean_sentences = pd.Series(sentences).str.replace("[^a-zA-Z]", " ")

# make alphabets lowercase
clean_sentences = [s.lower() for s in clean_sentences]
n=len(sentences)

nltk.download('stopwords')

from nltk.corpus import stopwords
stop_words = stopwords.words('english')


# function to remove stopwords
def remove_stopwords(sen):
    sen_new = " ".join([i for i in sen if i not in stop_words])
    return sen_new


# remove stopwords from the sentences
clean_sentences = [remove_stopwords(r.split()) for r in clean_sentences]



#Vector Representation of Sentences
# Extract word vectors
word_embeddings = {}
f = open('glove.6B.100d.txt', encoding='utf-8')
for line in f:
    values = line.split()
    word = values[0]
    coefs = np.asarray(values[1:], dtype='float32')
    word_embeddings[word] = coefs
f.close()


sentence_vectors = []
for i in clean_sentences:
  if len(i) != 0:
    v = sum([word_embeddings.get(w, np.zeros((100,))) for w in i.split()])/(len(i.split())+0.001)
  else:
    v = np.zeros((100,))
  sentence_vectors.append(v)
  
  
  
#Similarity Matrix Preparation
  # similarity matrix
sim_mat = np.zeros([len(sentences), len(sentences)])

from sklearn.metrics.pairwise import cosine_similarity

for i in range(len(sentences)):
  for j in range(len(sentences)):
    if i != j:
      sim_mat[i][j] = cosine_similarity(sentence_vectors[i].reshape(1,100), sentence_vectors[j].reshape(1,100))[0,0]
  
#Applying PageRank Algorithm
import networkx as nx

nx_graph = nx.from_numpy_array(sim_mat)
scores = nx.pagerank(nx_graph)

#Summary Extraction
ranked_sentences = sorted(((scores[i],s) for i,s in enumerate(sentences)), reverse=True)
# Extract top 10 sentences as the summary
ranked_sentences_proc=[]

for i in range(n//3):
  print(ranked_sentences[i][1])
  ranked_sentences_proc.append(ranked_sentences[i][1])
  
ranked_sentences_proc_df=pd.DataFrame(ranked_sentences_proc)
ranked_sentences_proc_df['Length']=ranked_sentences_proc_df[0].str.len()
ranked_sentences_proc_df['Rank']=ranked_sentences_proc_df['Length'].rank(ascending=False,method='average')

   
ranked_sentences_proc_df=ranked_sentences_proc_df.sort_values(by='Rank',ascending=True)
final=ranked_sentences_proc_df[0:7]
final = final.sort_index()
k=[]
for i in final.iloc[1:6,0]:
    k.append(i)
final_txt = ' '.join(k)
p = doc.add_paragraph(final_txt)
r = p.add_run()

##############To-do; Critical Items#####################
rawText=open('transcript.txt').read().strip()
##to-do and critical
rawText = rawText.lower()
rawText = rawText.replace('?','.')
rawText = rawText.split('.')



def check(sentence, words): 
    res = [all([k in s for k in words]) for s in sentence] 
    return [sentence[i] for i in range(0, len(res)) if res[i]] 
      
# Driver code 
#model = "When I'm on the courts or when I'm on the court playing, I'm a competitor and I want to beat every single person whether they're in the locker room or across the net.So I'm not the one to strike up a conversation about the weather and know that in the next few minutes I have to go and try to win a tennis match.I felt like the best weeks that I had to get to know players when I was playing were the Fed Cup weeks or the Olympic weeks, not necessarily during the tournaments."
#sentence_old = "Hello my name is alpha. I am going to help you. my agenda for today."
#checker = rawText.replace(".","','").strip()
#print(checker)
sentence = rawText
#sentence = ['Can you do this for me','make sure you do this','Can you also help','email','My important strategy is this ....','meeting','today','wrong','know','details','week' ,'day', 'help in this', 'my agenda', 'the conclusion..', 'brainstorm ...', 'collaborate ...','new update','what are the updates', 'think','feedback','everyone','input','result','method','point', 'important','agree','disagree','could','dont','do','recommend','should','how','approach','next','meeting','deadline','Working on this project', 'deadline is next weekend...', 'We must do ......', 'We could.....', 'Please check.....', 'Assure ....', 'Look into this....',] 
word_can = ['Can']
word_could = ['could']
word_should = ['should']
word_help = ['help']
word_agenda = ['agenda']
word_conclusion = ['conclusion']
word_today = ['today']
word_email =['email']
word_meeting  =['meeting']
word_wrong = ['wrong']
word_right = ['right']
word_know = ['know']
word_do = ['do']
word_dont = ['dont']
word_details = ['details']
word_week = ['week']
word_day = ['day']
word_brainstorm = ['brainstorm']
word_collaborate = ['collaborate']
word_update = ['update']
word_think = ['think']
word_feedback = ['feedback']
word_result = ['result']
word_method = ['method']
word_important = ['important']
word_point = ['point']
word_agree = ['agree']
word_disagree = ['disagree']
word_recommend = ['recommend']
word_dashboard = ['dashboard']
word_primarily = ['primarily']
word_performance =['performance']
word_phrase = ['important', 'strategy']
word_phrase_2 = ['make', 'sure']
word_phrase_3 = ['immediate', 'action']
word_immediate = ['immediate']
word_use = ['use']
word_year =["years"]
word_days =["days"]
word_want =["want"]


#note_can =print(check(sentence, word_can)) 
#note_help =print(check(sentence, word_help))
#note_agenda =print(check(sentence, word_agenda))
#note_conclusion =print(check(sentence, word_conclusion))
#note_dashboard =print(check(sentence, word_dashboard))

note_primarily =check(sentence, word_primarily)
note_year =(check(sentence, word_year))
note_days =(check(sentence, word_days))
note_want =(check(sentence, word_want))
note_immediate =(check(sentence, word_immediate))
#note_use =print(check(sentence, word_use))
#print(note_primarily)
#for i in note_primarily:
#    print(i.strip())
note_phrase_2 =(check(sentence, word_phrase_2))
note_phrase_3 =(check(sentence, word_phrase_3))


note_phrase_3 =(check(sentence, word_phrase_3))

#critical = print(note_immediate + note_phrase_2)
to_do= (note_immediate + note_phrase_2)

critical= (note_year,note_days,note_want)
#doc = docx.Document() 
#doc.add_heading('Meeting Notes - Global Executive Dashboard', 0) 
doc.add_heading('To-Do List', 1) 
for i in to_do:
    i = i.replace('\n','').strip().capitalize()
#    for a in  i:
    p = doc.add_paragraph(i, style='ListBullet')
    r = p.add_run()
doc.add_heading('Critical Items', 1) 
for i in critical:
#    i.strip()
#    i = i.capitalize()
#    count = count+1 
#    for a in  i:
#    print(str(count) +" "+i)
#    print(i)
    for a in i:
        a = a.replace('\n','').strip().capitalize()
        p = doc.add_paragraph(a, style='ListBullet')
        r = p.add_run()
WNL = nltk.WordNetLemmatizer()
 
# -----
##################WordCloud#########################################
def prepareStopWords():
 
    stopwordsList = []
 
    # Load default stop words and add a few more specific to my text.
    stopwordsList = stopwords.words('english')
    stopwordsList.append('dont')
    stopwordsList.append('didnt')
    stopwordsList.append('doesnt')
    stopwordsList.append('cant')
    stopwordsList.append('couldnt')
    stopwordsList.append('couldve')
    stopwordsList.append('im')
    stopwordsList.append('ive')
    stopwordsList.append('isnt')
    stopwordsList.append('theres')
    stopwordsList.append('wasnt')
    stopwordsList.append('wouldnt')
    stopwordsList.extend(['good','morning','a','immediate','require','like'])
    stopwordsList.extend(['also','action item','highlight area'])
    stopwordsList.extend(['give u','would','immediate attention','review meeting','good morning'])
    stopwordsList.extend(['Male Speaker 1', 'would like','information','sam','make sure','Kate'])
    return stopwordsList
 
# -----
 

# Open the file and read lines
# NOTE: You need to give finder.score_ngrams a sizable corpus to work with.
 

rawText=open('transcript.txt').read()


 
# Lowercase and tokenize

 
# Remove single quote early since it causes problems with the tokenizer.
# wasn't turns into 2 entries; was, n't.
rawText = rawText.replace("'", "").replace(".","")
 
tokens = nltk.word_tokenize(rawText)
text = nltk.Text(tokens)
 
# Load default stop words and add a few more.
stopWords = prepareStopWords()
 
# Remove extra chars and remove stop words.
text_content = [''.join(re.split("[ .,;:!?‘’``''@#$%^_&*()<>{}~\n\t\\\-]", word)) for word in text]
text_content = [word for word in text_content if word not in stopWords]
 
# After the punctuation above is removed it still leaves empty entries in the list.
# Remove any entries where the len is zero.
text_content = [s for s in text_content if len(s) != 0]
 
# Best to get the lemmas of each word to reduce the number of similar words
# on the word cloud. The default lemmatize method is noun, but this could be
# expanded.
# ex: The lemma of 'characters' is 'character'.
text_content = [WNL.lemmatize(t) for t in text_content]
 
# setup and score the bigrams using the raw frequency.
finder = BigramCollocationFinder.from_words(text_content)
bigram_measures = BigramAssocMeasures()
scored = finder.score_ngrams(bigram_measures.raw_freq)

# setup and score the trigram using the raw frequency.
#finder = nltk.collocations.TrigramCollocationFinder.from_words(text_content)
#trigram_measures = nltk.collocations.TrigramAssocMeasures()
#scored = finder.score_ngrams(trigram_measures.raw_freq)
# 
# By default finder.score_ngrams is sorted, however don't rely on this default behavior.
# Sort highest to lowest based on the score.
scoredList = sorted(scored, key=itemgetter(1), reverse=True)
 
# word_dict is the dictionary we'll use for the word cloud.
# Load dictionary with the FOR loop below.
# The dictionary will look like this with the bigram and the score from above.
# word_dict = {'bigram A': 0.000697411,
#             'bigram B': 0.000524882}
 
word_dict = {}
 
listLen = len(scoredList)
 
# Get the bigram and make a contiguous string for the dictionary key. 
# Set the key to the scored value. 
for i in range(listLen):
    word_dict[' '.join(scoredList[i][0])] = scoredList[i][1]


# -----
 
# Set word cloud params and instantiate the word cloud.
# The height and width only affect the output image file.
WC_height = 500
WC_width = 1000
WC_max_words = 20
 
wordCloud = WordCloud(max_words=WC_max_words, height=WC_height, width=WC_width,
                      background_color ='white')
 
wordCloud.generate_from_frequencies(word_dict)
 
plt.title('Most frequently occurring bigrams connected with an underscore_')
plt.imshow(wordCloud, interpolation='bilinear')
plt.axis("off")
plt.show()
wordCloud.to_file("WordCloud_Bigrams_frequent_words.png")
doc.add_heading('Most Frequent Words \n',1)
doc.add_picture("WordCloud_Bigrams_frequent_words.png",width=Inches(3.0), height=Inches(3.0))
doc.save('demo.docx')
